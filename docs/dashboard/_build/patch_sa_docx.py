# -*- coding: utf-8 -*-
"""「하남 SA 전략(초안)」 원본 docx 를 제자리에서 고친다.

앰플리튜드 수치를 xlsx 원본·대시보드와 일치시킨다. 서식은 건드리지 않고
런(run) 텍스트만 갈아끼운다 — 표 모양·열 수·색은 원본 그대로 유지.
표 세 개는 행 수가 그대로라 셀 값만 덮어쓰면 된다.
"""
import io
import sys
from docx import Document

SRC = r'C:\Users\user\Downloads\하남 SA 전략(초안) (1).docx'
OUT = r'C:\Users\user\Desktop\GMM\0709_claude_하남_웹페이지제작\docs\하남 SA 전략(수정본).docx'

d = Document(SRC)
done = []
miss = []


def set_text(par, new):
    """첫 런의 서식을 유지한 채 문단 전체 텍스트를 교체한다."""
    if not par.runs:
        par.add_run(new)
        return
    par.runs[0].text = new
    for r in par.runs[1:]:
        r.text = ''


def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    yield p


def sub(needle, new, label, whole=True):
    """needle 을 포함하는 문단을 찾아 교체. whole=False 면 부분 치환."""
    hit = 0
    for p in all_paragraphs(d):
        txt = p.text
        if needle in txt:
            set_text(p, new if whole else txt.replace(needle, new))
            hit += 1
    (done if hit else miss).append(f'{label}  ({hit}건)')
    return hit


def fill_table(match_cell, rows, label):
    """첫 행 첫 칸이 match_cell 로 시작하는 표를 찾아 rows 로 덮어쓴다."""
    for t in d.tables:
        head = t.rows[0].cells[0].text.strip()
        if head != match_cell:
            continue
        if len(t.rows) != len(rows):
            miss.append(f'{label} 행 수 불일치 {len(t.rows)} vs {len(rows)}')
            return False
        for tr, vals in zip(t.rows, rows):
            cells = tr.cells
            if len(cells) != len(vals):
                miss.append(f'{label} 열 수 불일치 {len(cells)} vs {len(vals)}')
                return False
            for c, v in zip(cells, vals):
                ps = c.paragraphs
                set_text(ps[0], v)
                for extra in ps[1:]:
                    set_text(extra, '')
        done.append(f'{label}  ({len(rows)}행)')
        return True
    miss.append(label + ' 표 못 찾음')
    return False


# ══ 장표 04 ═══════════════════════════════════════════════════════
sub('퍼널 도표 — 클릭 17,828 → 도착 8,374 → CTA 108',
    '퍼널 도표 — 고유 클릭 14,570 → 도착 8,084 → CTA 110 → 예약 유도 90',
    '04 ① 퍼널 도표')

sub('클릭 17,828 → 도착 8,374 · 노출의 64%가 릴스',
    '고유 클릭 14,570 → 도착 8,084 (55.5%) · 노출의 64%가 릴스',
    '04 ② 표 1행')

# 2행 — 체류·Dead Click·언어 전환 셋 다 근거가 없어 통째로 교체
sub('체류 13초 · Dead Click 74 · 메뉴 본 사람의 53%가 언어 전환',
    '방문자의 95.1%가 아무것도 누르지 않고 이탈 · 예약 CTA의 85.9%가 첫 화면·플로팅에 집중',
    '04 ② 표 2행 (데이터)')
sub('첫 화면에 예약 버튼 고정 + 언어 자동 감지',
    '첫 화면에 예약 버튼 고정',
    '04 ② 표 2행 (실행)')

sub('예약 행동의 44.4%가 전화·길찾기로 이탈 → 추적 불가',
    '예약 행동의 45.9%가 전화·길찾기로 이탈 → 추적 불가',
    '04 ② 표 3행')

# 대본은 4개 단락으로 나뉘어 있다. 단락 구분을 그대로 두고 각각 갈아끼운다.
SCRIPT = [
    '저희 전략은 아이디어가 아니라 이 데이터에서 그대로 나왔습니다.',

    '광고를 클릭한 1만 4천 5백 명 중 저희 페이지를 실제로 본 사람은 8천 명이었습니다. '
    '45%가 도착 전에 사라졌고, 노출의 64%가 릴스였습니다. '
    '스크롤하다 잘못 누른 클릭이라는 뜻입니다. 그래서 의도를 갖고 누르는 검색으로 옮깁니다.',

    '도착한 사람도 95%는 아무것도 누르지 않고 나갔습니다. '
    '반면 예약 버튼을 누른 사람의 86%는 첫 화면과 플로팅 버튼에서 눌렀습니다. '
    '버튼은 제 역할을 했다는 뜻입니다. 그래서 첫 화면에 예약 버튼을 고정합니다.',

    '그리고 예약 행동의 45.9%가 전화와 길찾기로 페이지를 빠져나가 그 뒤를 추적할 수 없었습니다. '
    '광고로 데려와도 예약이 됐는지 알 수가 없습니다. '
    '그래서 광고가 도착할 «예약 전용 페이지»가 필요합니다. '
    '예약이 한 곳에서 끝나야 광고에서 방문까지가 하나의 숫자로 이어집니다.',
]
_hit = False
for _t in d.tables:
    for _r in _t.rows:
        for _c in _r.cells:
            if '저희 전략은 아이디어가' not in _c.text:
                continue
            _ps = [x for x in _c.paragraphs if x.text.strip()]
            if len(_ps) != len(SCRIPT):
                miss.append(f'04 ③ 대본 단락 수 불일치 {len(_ps)} vs {len(SCRIPT)}')
            for _p, _new in zip(_ps, SCRIPT):
                set_text(_p, _new)
            _hit = True
(done if _hit else miss).append('04 ③ 읽을 대본  (4단락)')

sub('44.4%의 근거', '45.9%의 근거', '04 ④ 근거 제목', whole=False)

fill_table('행동', [
    ['행동', '건수', '비중', '추적'],
    ['메시지 예약', '61', '45.9%', '가능'],
    ['전화 걸기', '45', '33.8%', '불가'],
    ['길찾기', '16', '12.0%', '불가'],
    ['메뉴 보기', '11', '8.3%', '가능'],
    ['합계', '133', '100%', '45.9% 추적 불가'],
], '04 ④ 근거 표')

sub('cta_type별로 분해한 값 (Uniques, 2026.7.20–8.2). 정확.',
    'cta_type별로 분해한 값 (이벤트 기준, 2026.7.20–8.2). 정확. '
    '사용자 기준은 110명이며, 한 사람이 전화와 길찾기를 모두 누르면 중복 집계되어 비중 계산에 쓸 수 없다.',
    '04 ④ 근거 각주', whole=False)

# ══ 장표 05 ═══════════════════════════════════════════════════════
sub('예약 행동 108건 중 71건(65.7%)이 이 소재 하나에서 나왔고, 전환당 비용도 58,437동으로 전체 평균 81,634동보다 낮음.',
    '예약 유도 90명 중 64명(71.1%)이 이 소재 하나에서 나왔고, 유도당 비용도 64,828동으로 전체 평균 97,961동보다 낮음.',
    '05 마지막 각주', whole=False)

# ══ 부록 ② ════════════════════════════════════════════════════════
fill_table('지표', [
    ['지표', '값', '원천 · 산출'],
    ['노출', '460,195', '메타 광고 관리자 (정확)'],
    ['전체 링크클릭', '17,828', '메타 — CTR·CPC 산출 기준'],
    ['고유 링크클릭', '14,570', '메타 — 도착률 분모'],
    ['지출', '8,816,522 VND', '메타 (예산 소진)'],
    ['랜딩 도착', '8,084  (도착률 55.5%)', '엠플리튜드 · 고유 device_id'],
    ['도착 전 이탈', '6,486  (44.5%)', '계산값 = 14,570 − 8,084'],
    ['CTA 클릭', '110명 (133건)  ·  1.36%', '엠플리튜드 「CTA Clicked」'],
    ['예약 유도', '90명 (106건)  ·  1.11%', 'cta_type ∈ {message_book, call_phone}'],
    ['무행동 이탈', '95.1%', '아무 버튼도 누르지 않은 방문자'],
    ['유도당 비용', '97,961 VND  (약 5,267원)', '지출 ÷ 예약 유도 90명 · 환율 18.6'],
    ['릴스 노출 비중', '63.9%  (293,874 ÷ 460,195)', '메타 지면 데이터'],
    ['전환 소요 시간 중앙값', '12.5초', '예약 CTA를 누른 94명 한정 — 체류시간 아님'],
], '부록② 숫자 표')

sub('«도착 전 이탈»과 «도착 후 이탈»은 원본에 없는 계산값입니다.',
    '«도착 전 이탈»은 원본에 없는 계산값입니다. 도착률은 고유 클릭 기준으로 산출했습니다 — '
    '방문자가 사람 수이므로, 같은 사람의 반복 클릭을 포함한 전체 클릭과는 단위가 맞지 않습니다.',
    '부록② 각주', whole=False)

fill_table('소재', [
    ['소재', '노출', '고유 클릭', '예약 유도', '유도당 비용(VND)'],
    ['비즈니스 × 맛 (영상)', '199,294', '7,955', '64  (71.1%)', '64,828'],
    ['가족 × 맛 (영상)', '141,222', '3,881', '11', '238,690'],
    ['가족 × 서비스 (영상)', '65,218', '2,620', '12', '124,832'],
    ['가족 × 서비스 (이미지)', '37,899', '273', '2  (N작음)', '116,188'],
    ['비즈니스 × 서비스 (영상)', '5,022', '168', '0  (N작음)', '—'],
    ['비즈니스 × 서비스 (캐러셀)', '11,540', '227', '0  (N작음)', '—'],
], '부록② 소재별 표')

sub('「비즈니스 × 서비스(영상)」의 CPA가 더 낮아 보이지만',
    '방문 N < 300 소재(가족×서비스 이미지 120 · 비즈니스×서비스 영상 122 · 캐러셀 48)는 표본이 작아 '
    '비율 비교에서 제외했습니다. 「비즈니스 × 맛」은 물량과 성과를 동시에 가진 유일한 소재이며, '
    '예약 유도 90명 중 64명을 이 소재 하나가 만들었습니다. '
    '고유 클릭은 사람 수라 중복이 제거되어 소재별 합(15,124)이 전체(14,570)와 다릅니다.',
    '부록② 소재 각주')

d.save(OUT)

print('■ 적용됨')
for x in done:
    print('   OK  ', x)
if miss:
    print('■ 확인 필요')
    for x in miss:
        print('   !!  ', x)
print('\nsaved ->', OUT)
