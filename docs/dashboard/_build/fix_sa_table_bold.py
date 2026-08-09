# -*- coding: utf-8 -*-
"""표 굵기 정리.

행 텍스트만 갈아끼우면서 원본 런의 굵기가 그대로 남아, 새로 들어간 행과
어긋났다. 표별로 규칙을 하나씩 정해 다시 칠한다.
"""
from docx import Document

P = r'C:\Users\user\Desktop\GMM\0709_claude_하남_웹페이지제작\docs\하남 SA 전략(수정본).docx'
OUTP = 'C:\\Users\\user\\AppData\\Local\\Temp\\claude\\C--Users-user-Desktop-GMM-0709-claude----------\\15c39293-391c-48e6-a986-c2a68cdec763\\scratchpad\\하남 SA 전략(수정본).docx'
d = Document(P)


def paint(cell, bold):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold


for t in d.tables:
    head = t.rows[0].cells[0].text.strip()

    # ① 45.9% 근거 — 추적 불가인 두 행과 합계만 굵게
    if head == '행동':
        for i, r in enumerate(t.rows):
            if i == 0:
                continue
            key = r.cells[0].text.strip()
            b = key in ('전화 걸기', '길찾기', '합계')
            for c in r.cells:
                paint(c, b)

    # ② 부록② 숫자 표 — 지표명과 값은 굵게, 출처는 보통. 행마다 다르면 산만하다.
    elif head == '지표':
        for i, r in enumerate(t.rows):
            if i == 0:
                continue
            cs = r.cells
            paint(cs[0], True)
            paint(cs[1], True)
            paint(cs[2], False)

    # ③ 소재별 표 — 기준 소재 한 줄만 굵게 (원본 의도 유지)
    elif head == '소재':
        for i, r in enumerate(t.rows):
            if i == 0:
                continue
            b = r.cells[0].text.strip().startswith('비즈니스 × 맛')
            for c in r.cells:
                paint(c, b)

d.save(OUTP)
print('굵기 정리 완료')

# 확인
d2 = Document(OUTP)
for t in d2.tables:
    h = t.rows[0].cells[0].text.strip()
    if h not in ('행동', '지표', '소재'):
        continue
    print('===', h)
    for r in t.rows:
        m = []
        for c in r.cells:
            bs = [bool(x.bold) for p in c.paragraphs for x in p.runs if x.text.strip()]
            m.append(('B:' if any(bs) else '  ') + c.text.strip()[:20])
        print('  ', ' | '.join(m))
